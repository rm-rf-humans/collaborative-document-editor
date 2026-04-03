from __future__ import annotations

import html
import re
from io import BytesIO
from typing import Literal

from docx import Document as DocxDocument
from fastapi import FastAPI
from fastapi.responses import Response
from pydantic import BaseModel, Field
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, Preformatted, SimpleDocTemplate, Spacer


class ExportedAiInteraction(BaseModel):
    feature: str
    status: str
    initiatedBy: str
    createdAt: str
    requestedVersion: int
    sourceText: str
    suggestedText: str
    targetLanguage: str | None = None


class ExportRequest(BaseModel):
    format: Literal["markdown", "pdf", "docx"]
    title: str = Field(min_length=1, max_length=200)
    content: str
    version: int = Field(ge=1)
    generatedAt: str
    includeAiAppendix: bool = True
    aiInteractions: list[ExportedAiInteraction] = Field(default_factory=list)


app = FastAPI(title="CollabWrite Export API", version="1.0.0")


def sanitize_filename(value: str) -> str:
    normalized = re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-")
    return normalized[:64] or "document"


def split_content_lines(content: str) -> list[str]:
    return content.splitlines() or [""]


def markdown_quote_block(value: str) -> str:
    return "\n".join(f"> {line}" for line in (value.splitlines() or [""]))


def build_markdown(request: ExportRequest) -> bytes:
    sections = [
        f"# {request.title}",
        "",
        f"- Snapshot version: v{request.version}",
        f"- Generated at: {request.generatedAt}",
        "",
        request.content.rstrip(),
    ]

    if request.includeAiAppendix:
        sections.extend(["", "## AI Suggestion History"])
        if not request.aiInteractions:
            sections.extend(["", "_No AI interactions recorded for this snapshot._"])
        else:
            for index, interaction in enumerate(request.aiInteractions, start=1):
                sections.extend(
                    [
                        "",
                        f"### {index}. {interaction.feature} ({interaction.status})",
                        "",
                        f"- Initiated by: {interaction.initiatedBy}",
                        f"- Requested version: v{interaction.requestedVersion}",
                        f"- Created at: {interaction.createdAt}",
                    ]
                )
                if interaction.targetLanguage:
                    sections.append(f"- Target language: {interaction.targetLanguage}")
                sections.extend(
                    [
                        "",
                        "Source text:",
                        markdown_quote_block(interaction.sourceText),
                        "",
                        "Suggested text:",
                        markdown_quote_block(interaction.suggestedText or "_No suggestion text was stored._"),
                    ]
                )

    return "\n".join(sections).encode("utf-8")


def append_docx_content(doc: DocxDocument, content: str) -> None:
    for line in split_content_lines(content):
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")


def build_docx(request: ExportRequest) -> bytes:
    doc = DocxDocument()
    doc.add_heading(request.title, level=0)
    doc.add_paragraph(f"Snapshot version: v{request.version}")
    doc.add_paragraph(f"Generated at: {request.generatedAt}")
    append_docx_content(doc, request.content)

    if request.includeAiAppendix:
        doc.add_page_break()
        doc.add_heading("AI Suggestion History", level=1)
        if not request.aiInteractions:
            doc.add_paragraph("No AI interactions recorded for this snapshot.")
        else:
            for index, interaction in enumerate(request.aiInteractions, start=1):
                doc.add_heading(f"{index}. {interaction.feature} ({interaction.status})", level=2)
                doc.add_paragraph(f"Initiated by: {interaction.initiatedBy}")
                doc.add_paragraph(f"Requested version: v{interaction.requestedVersion}")
                doc.add_paragraph(f"Created at: {interaction.createdAt}")
                if interaction.targetLanguage:
                    doc.add_paragraph(f"Target language: {interaction.targetLanguage}")
                doc.add_paragraph("Source text:")
                doc.add_paragraph(interaction.sourceText or "(empty)")
                doc.add_paragraph("Suggested text:")
                doc.add_paragraph(interaction.suggestedText or "(empty)")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def pdf_paragraph(text: str, style_name: str, styles) -> Paragraph:
    return Paragraph(html.escape(text), styles[style_name])


def append_pdf_content(story: list, content: str, styles) -> None:
    for line in split_content_lines(content):
        stripped = line.strip()
        if stripped.startswith("# "):
            story.append(pdf_paragraph(stripped[2:].strip(), "Heading1", styles))
        elif stripped.startswith("## "):
            story.append(pdf_paragraph(stripped[3:].strip(), "Heading2", styles))
        elif stripped:
            story.append(pdf_paragraph(stripped, "BodyText", styles))
        else:
            story.append(Spacer(1, 8))


def build_pdf(request: ExportRequest) -> bytes:
    buffer = BytesIO()
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="Metadata", parent=styles["BodyText"], textColor="#5A5A5A"))

    story: list = [
        pdf_paragraph(request.title, "Title", styles),
        pdf_paragraph(f"Snapshot version: v{request.version}", "Metadata", styles),
        pdf_paragraph(f"Generated at: {request.generatedAt}", "Metadata", styles),
        Spacer(1, 12),
    ]
    append_pdf_content(story, request.content, styles)

    if request.includeAiAppendix:
        story.extend([Spacer(1, 18), pdf_paragraph("AI Suggestion History", "Heading1", styles)])
        if not request.aiInteractions:
            story.append(pdf_paragraph("No AI interactions recorded for this snapshot.", "BodyText", styles))
        else:
            for index, interaction in enumerate(request.aiInteractions, start=1):
                story.extend(
                    [
                        Spacer(1, 12),
                        pdf_paragraph(f"{index}. {interaction.feature} ({interaction.status})", "Heading2", styles),
                        pdf_paragraph(f"Initiated by: {interaction.initiatedBy}", "Metadata", styles),
                        pdf_paragraph(f"Requested version: v{interaction.requestedVersion}", "Metadata", styles),
                        pdf_paragraph(f"Created at: {interaction.createdAt}", "Metadata", styles),
                    ]
                )
                if interaction.targetLanguage:
                    story.append(pdf_paragraph(f"Target language: {interaction.targetLanguage}", "Metadata", styles))
                story.extend(
                    [
                        Spacer(1, 6),
                        pdf_paragraph("Source text", "Heading3", styles),
                        Preformatted(interaction.sourceText or "(empty)", styles["Code"]),
                        Spacer(1, 6),
                        pdf_paragraph("Suggested text", "Heading3", styles),
                        Preformatted(interaction.suggestedText or "(empty)", styles["Code"]),
                    ]
                )

    document = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=54, rightMargin=54, topMargin=54, bottomMargin=54)
    document.build(story)
    return buffer.getvalue()


def render_export(request: ExportRequest) -> tuple[bytes, str, str]:
    filename = f"{sanitize_filename(request.title)}-v{request.version}"
    if request.format == "markdown":
        return build_markdown(request), "text/markdown; charset=utf-8", f"{filename}.md"
    if request.format == "docx":
        return build_docx(
            request
        ), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", f"{filename}.docx"
    return build_pdf(request), "application/pdf", f"{filename}.pdf"


@app.get("/health")
def healthcheck() -> dict[str, str]:
    return {"status": "ok"}


@app.post("/render")
def render(request: ExportRequest) -> Response:
    payload, content_type, filename = render_export(request)
    return Response(
        content=payload,
        media_type=content_type,
        headers={"content-disposition": f'attachment; filename="{filename}"'},
    )
